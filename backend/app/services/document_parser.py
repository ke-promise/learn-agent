from __future__ import annotations

from io import BytesIO
from pathlib import Path
from typing import Any

from fastapi import UploadFile

from app.core.config import settings


class DocumentParser:
    """文档解析器。

    它的职责是把不同来源的文件统一转成纯文本，供后续切块和向量化使用。
    """

    def __init__(self) -> None:
        self._ocr_engine: Any | None = None

    async def parse(self, file: UploadFile | None, content: str) -> tuple[str, str | None, bytes | None]:
        """统一处理“纯文本输入”和“文件上传输入”两种场景。"""
        if file is None:
            return content.strip(), None, None
        raw = await file.read()
        filename = file.filename or 'upload'
        return self.extract_text(raw, filename), filename, raw

    def extract_text(self, raw: bytes, filename: str) -> str:
        """根据文件后缀选择解析策略。"""
        suffix = Path(filename).suffix.lower()
        if suffix == '.pdf':
            text = self._extract_pdf_text(raw)
            if text.strip():
                return text
            if settings.ocr_enabled:
                return self._ocr_pdf(raw)
            return text
        if suffix == '.docx':
            return self._extract_docx_text(raw)
        if suffix in {'.png', '.jpg', '.jpeg', '.bmp'} and settings.ocr_enabled:
            return self._ocr_image(raw)
        return raw.decode('utf-8', errors='ignore')

    def _extract_pdf_text(self, raw: bytes) -> str:
        """优先尝试从 PDF 中直接抽取文本。"""
        try:
            from pypdf import PdfReader
        except ImportError:
            return raw.decode('utf-8', errors='ignore')
        reader = PdfReader(BytesIO(raw))
        pages = []
        for page in reader.pages:
            text = page.extract_text() or ''
            if text.strip():
                pages.append(text.strip())
        return '\n\n'.join(pages).strip()

    def _extract_docx_text(self, raw: bytes) -> str:
        """提取 DOCX 段落和表格内容。"""
        try:
            from docx import Document
        except ImportError:
            return raw.decode('utf-8', errors='ignore')
        document = Document(BytesIO(raw))
        sections: list[str] = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                sections.append(text)
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    sections.append(' | '.join(cells))
        return '\n'.join(sections).strip()

    def _ocr_pdf(self, raw: bytes) -> str:
        """扫描版 PDF 无法直接抽文本时，逐页转图后走 OCR。"""
        try:
            import fitz
        except ImportError:
            return ''
        pages: list[str] = []
        document = fitz.open(stream=raw, filetype='pdf')
        for index, page in enumerate(document):
            if index >= settings.ocr_max_pdf_pages:
                break
            pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2))
            pages.append(self._ocr_image(pixmap.tobytes('png')))
        return '\n\n'.join(text for text in pages if text.strip()).strip()

    def _ocr_image(self, raw: bytes) -> str:
        """对图片直接做 OCR。"""
        try:
            import numpy as np
            from PIL import Image
        except ImportError:
            return ''
        engine = self._get_ocr_engine()
        if engine is None:
            return ''
        image = Image.open(BytesIO(raw)).convert('RGB')
        result, _ = engine(np.array(image))
        texts: list[str] = []
        for item in result or []:
            if not isinstance(item, (list, tuple)) or len(item) < 2:
                continue
            candidate = item[1]
            if isinstance(candidate, (list, tuple)) and candidate:
                candidate = candidate[0]
            if candidate:
                texts.append(str(candidate).strip())
        return '\n'.join(text for text in texts if text)

    def _get_ocr_engine(self) -> Any | None:
        """按需懒加载 OCR 引擎，避免启动时就加载大依赖。"""
        if self._ocr_engine is not None:
            return self._ocr_engine
        try:
            from rapidocr_onnxruntime import RapidOCR
        except ImportError:
            return None
        self._ocr_engine = RapidOCR()
        return self._ocr_engine